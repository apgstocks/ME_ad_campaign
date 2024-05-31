import time
import pywhatkit
from docx import Document
import pyautogui
from pynput.keyboard import Key,Controller
import pandas as pd
from locator import *
keyboard=Controller()
def send_whatsapp_message(msg,phone,img):
    try:
        for ph in phone:
           
            ph=str('+91'+str(ph))
            print(ph)
            try:
                pywhatkit.sendwhats_image(ph,img,msg,tab_close=True)
                time.sleep(2)
                pyautogui.click()
                keyboard.press(Key.enter)
                keyboard.release(Key.enter)
            except Exception as e:
                print('some error occured while sending image')
                print(e)
            """
            pywhatkit.sendwhatmsg_instantly(phone_no=ph,message=msg,tab_close=True)
            time.sleep(2)
            pyautogui.click()
            time.sleep(2)
            keyboard.press(Key.enter)
            keyboard.release(Key.enter)
            """
            print(f"Message sent to {ph}")
    except Exception as e:
        print(str(e))

def getText(filename):
    doc = Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

if __name__=='__main__':
    #phone_numbers=["+918056944193","+919486076039","+919487559628"]
    #message="Welcome to Meat and Eat"
    image_path=image
    #send_whatsapp_message(message,phone_numbers)
    #pd.read_excel("")
    try:
        customer=pd.read_excel(excel,sheet_name= "Sheet 1", engine='openpyxl')
        
        msg=getText(message)
        #Document(r'/Users/apsara/Documents/ME_message.docx')
        #print(msg)
        #print(customer.head())
        #print(customer.columns)
        send_whatsapp_message(msg,customer['Phone_numbers'],image_path)
    except Exception as e:
        print(str(e))

            